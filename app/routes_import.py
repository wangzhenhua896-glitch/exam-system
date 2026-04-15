"""导入/导出 API"""
import json
from flask import request, jsonify, session, Response
from app.api_shared import api_bp
from app.api_shared import _session_subject, _check_subject_access
from app.models.db_models import get_questions, add_question
from loguru import logger


@api_bp.route('/export-rubric-scripts', methods=['GET'])
def export_rubric_scripts():
    """按科目导出评分脚本 — 非 admin 强制用 session.subject"""
    from flask import Response
    import io
    subject = _session_subject() or request.args.get('subject', '').strip()
    fmt = request.args.get('format', 'md').strip().lower()
    if not subject:
        return jsonify({'success': False, 'error': '请指定科目参数'}), 400
    questions = get_questions(subject)

    # 过滤有评分脚本的题目
    questions_with_script = [q for q in questions if (q.get('rubric_script') or '').strip()]
    if not questions_with_script:
        return jsonify({'success': False, 'error': '没有找到包含评分脚本的题目'}), 404

    # 按科目分组
    grouped = {}
    for q in questions_with_script:
        subj = q.get('subject', '未分类')
        grouped.setdefault(subj, []).append(q)

    subject_labels = {
        'politics': '思想政治', 'chinese': '语文', 'english': '英语',
        'math': '数学', 'history': '历史', 'geography': '地理',
        'physics': '物理', 'chemistry': '化学', 'biology': '生物',
    }

    if fmt == 'docx':
        # 生成 Word 文档
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()

        # 标题
        title_para = doc.add_heading('评分脚本导出', level=0)
        # 元信息
        meta = doc.add_paragraph()
        meta_run = meta.add_run(f'导出时间：{__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M")}　　共 {len(questions_with_script)} 道题')
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = __import__('docx.shared', fromlist=['RGBColor']).RGBColor(0x90, 0x93, 0x99)
        doc.add_paragraph('')

        for subj, qs in grouped.items():
            label = subject_labels.get(subj, subj)
            doc.add_heading(label, level=1)

            for q in qs:
                title = q.get('title') or q.get('content', '')[:30]
                max_score = q.get('max_score', '?')
                doc.add_heading(f'{title}（{max_score}分）', level=2)

                # 题目内容
                content = q.get('content', '').strip()
                if content:
                    p = doc.add_paragraph()
                    run = p.add_run('题目：')
                    run.bold = True
                    p.add_run('\n' + content)

                # 标准答案
                standard_answer = q.get('standard_answer', '').strip()
                if standard_answer:
                    p = doc.add_paragraph()
                    run = p.add_run('标准答案：')
                    run.bold = True
                    p.add_run('\n' + standard_answer)

                # 评分脚本
                rubric_script = q.get('rubric_script', '').strip()
                if rubric_script:
                    p = doc.add_paragraph()
                    run = p.add_run('评分脚本：')
                    run.bold = True
                    # 逐行添加，保留格式
                    for line in rubric_script.split('\n'):
                        p.add_run('\n' + line)

                doc.add_paragraph('')  # 题目间空行

        buf = io.BytesIO()
        doc.save(buf)
        doc_bytes = buf.getvalue()

        from urllib.parse import quote
        filename = f'评分脚本_{subject}.docx'
        filename_encoded = quote(filename)

        return Response(
            doc_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f"attachment; filename*=UTF-8''{filename_encoded}",
            }
        )

    # 默认：Markdown 格式
    lines = []
    lines.append('# 评分脚本导出')
    lines.append('')
    lines.append(f'> 导出时间：{__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M")}')
    lines.append(f'> 共 {len(questions_with_script)} 道题')
    lines.append('')
    lines.append('---')
    lines.append('')

    for subj, qs in grouped.items():
        label = subject_labels.get(subj, subj)
        lines.append(f'## {label}')
        lines.append('')

        for q in qs:
            title = q.get('title') or q.get('content', '')[:30]
            max_score = q.get('max_score', '?')
            lines.append(f'### {title}（{max_score}分）')
            lines.append('')

            content = q.get('content', '').strip()
            if content:
                lines.append('**题目：**')
                lines.append('')
                lines.append(content)
                lines.append('')

            standard_answer = q.get('standard_answer', '').strip()
            if standard_answer:
                lines.append('**标准答案：**')
                lines.append('')
                lines.append(standard_answer)
                lines.append('')

            rubric_script = q.get('rubric_script', '').strip()
            if rubric_script:
                lines.append('**评分脚本：**')
                lines.append('')
                lines.append(rubric_script)
                lines.append('')

            lines.append('---')
            lines.append('')

    md_content = '\n'.join(lines)

    from urllib.parse import quote
    filename = f'评分脚本_{subject}.md'
    filename_encoded = quote(filename)

    return Response(
        md_content,
        mimetype='text/markdown',
        headers={
            'Content-Disposition': f"attachment; filename*=UTF-8''{filename_encoded}",
            'Content-Type': 'text/markdown; charset=utf-8',
        }
    )


@api_bp.route('/import-questions/preview', methods=['POST'])
def import_questions_preview():
    """预览文件：返回各 sheet/表名、列名、行数、前几行数据、自动列映射"""
    import pandas as pd

    if 'file' not in request.files:
        return jsonify({'success': False, 'error': '未上传文件'}), 400

    file = request.files['file']
    fname = file.filename.lower()
    if not (fname.endswith(('.xlsx', '.xls', '.csv')) or fname.endswith('.docx')):
        return jsonify({'success': False, 'error': '仅支持 .xlsx/.xls/.csv/.docx 文件'}), 400

    # 数据库字段定义
    DB_FIELDS = [
        {'key': 'question_number', 'label': '题号', 'required': False, 'hints': ('题号', '编号', 'number', 'question_number', 'no', '序号')},
        {'key': 'content', 'label': '题目内容', 'required': True, 'hints': ('题目', '题目内容', 'content', 'question', '题干')},
        {'key': 'title', 'label': '标题', 'required': False, 'hints': ('标题', 'title', '题目标题')},
        {'key': 'standard_answer', 'label': '标准答案', 'required': False, 'hints': ('标准答案', '答案', 'answer', 'standard_answer')},
        {'key': 'max_score', 'label': '分值', 'required': False, 'hints': ('满分', '分值', 'score', 'max_score')},
        {'key': 'difficulty', 'label': '难度', 'required': False, 'hints': ('难度', 'difficulty', '难易')},
        {'key': 'exam_name', 'label': '试卷名称', 'required': False, 'hints': ('试卷', '试卷名称', 'exam', 'exam_name', '考试')},
        {'key': 'rubric_points', 'label': '评分要点', 'required': False, 'hints': ('得分点', '评分要点', 'points', 'rubric_points')},
        {'key': 'rubric_rules', 'label': '评分规则', 'required': False, 'hints': ('评分规则', 'rules', 'rubric_rules')},
    ]

    def guess_mapping(columns):
        """根据列名自动猜测映射关系"""
        mapping = {}
        mapped_cols = set()
        for field in DB_FIELDS:
            for col in columns:
                cl = str(col).strip().lower()
                if cl in [h.lower() for h in field['hints']]:
                    mapping[field['key']] = col
                    mapped_cols.add(col)
                    break
        # 找出未映射的列
        unmapped = [c for c in columns if c not in mapped_cols]
        return mapping, unmapped

    try:
        if file.filename.lower().endswith('.docx'):
            # Word 文件：解析其中的表格
            from docx import Document
            import io
            doc = Document(io.BytesIO(file.read()))
            sheets = []
            for t_idx, table in enumerate(doc.tables):
                rows_data = []
                max_cols = 0
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    max_cols = max(max_cols, len(cells))
                    rows_data.append(cells)
                # 补齐列数不一致的行（合并单元格导致）
                rows_data = [r + [''] * (max_cols - len(r)) for r in rows_data]
                if len(rows_data) < 2:
                    continue  # 跳过空表或只有表头的表
                columns = rows_data[0]
                data_rows = rows_data[1:]
                df = pd.DataFrame(data_rows, columns=columns)
                preview_rows = df.head(5).fillna('').to_dict(orient='records')
                auto_map, unmapped = guess_mapping(df.columns)
                table_name = f'表格{t_idx + 1}'
                sheets.append({'name': table_name, 'columns': list(df.columns), 'rows': len(df), 'preview': preview_rows, 'mapping': auto_map, 'unmapped_columns': unmapped})
            if not sheets:
                return jsonify({'success': False, 'error': 'Word 文件中未找到有效表格（至少需要表头+1行数据）'}), 400
        elif file.filename.endswith('.csv'):
            df = pd.read_csv(file)
            preview_rows = df.head(5).fillna('').to_dict(orient='records')
            auto_map, unmapped = guess_mapping(df.columns)
            sheets = [{'name': 'Sheet1', 'columns': list(df.columns), 'rows': len(df), 'preview': preview_rows, 'mapping': auto_map, 'unmapped_columns': unmapped}]
        else:
            xls = pd.ExcelFile(file)
            sheets = []
            for name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=name)
                preview_rows = df.head(5).fillna('').to_dict(orient='records')
                auto_map, unmapped = guess_mapping(df.columns)
                sheets.append({'name': name, 'columns': list(df.columns), 'rows': len(df), 'preview': preview_rows, 'mapping': auto_map, 'unmapped_columns': unmapped})
    except Exception as e:
        return jsonify({'success': False, 'error': f'读取文件失败：{str(e)}'}), 400

    return jsonify({'success': True, 'data': {
        'filename': file.filename,
        'sheets': sheets,
        'db_fields': DB_FIELDS,
    }})


@api_bp.route('/import-questions', methods=['POST'])
def import_questions():
    """从文件导入题目（支持 Excel/CSV/Word，支持指定 sheet 和科目）"""
    import pandas as pd

    if 'file' not in request.files:
        return jsonify({'success': False, 'error': '未上传文件'}), 400

    file = request.files['file']
    fname = file.filename.lower()
    if not (fname.endswith(('.xlsx', '.xls', '.csv')) or fname.endswith('.docx')):
        return jsonify({'success': False, 'error': '仅支持 .xlsx/.xls/.csv/.docx 文件'}), 400

    # 支持指定要导入的 sheet、对应科目和列映射
    import json as _json
    sheet_configs = request.form.get('sheet_configs')  # JSON: [{"sheet":"Sheet1","subject":"politics","col_map":{"content":"题目"},"extra_fields":{}},...]
    if sheet_configs:
        sheet_configs = _json.loads(sheet_configs)
    else:
        sheet_configs = None

    try:
        if file.filename.lower().endswith('.docx'):
            # Word 文件：解析其中的表格
            from docx import Document
            import io
            doc = Document(io.BytesIO(file.read()))
            dfs = []
            for t_idx, table in enumerate(doc.tables):
                rows_data = []
                max_cols = 0
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    max_cols = max(max_cols, len(cells))
                    rows_data.append(cells)
                rows_data = [r + [''] * (max_cols - len(r)) for r in rows_data]
                if len(rows_data) < 2:
                    continue
                columns = rows_data[0]
                data_rows = rows_data[1:]
                df = pd.DataFrame(data_rows, columns=columns)
                table_name = f'表格{t_idx + 1}'
                if sheet_configs and t_idx < len(sheet_configs):
                    dfs.append((sheet_configs[t_idx].get('sheet', table_name), df))
                else:
                    dfs.append((table_name, df))
            if not dfs:
                return jsonify({'success': False, 'error': 'Word 文件中未找到有效表格'}), 400
            xls = None
        elif file.filename.endswith('.csv'):
            xls = None
            dfs = [('Sheet1', pd.read_csv(file))]
        else:
            xls = pd.ExcelFile(file)
            if sheet_configs:
                dfs = [(sc['sheet'], pd.read_excel(xls, sheet_name=sc['sheet'])) for sc in sheet_configs]
            else:
                dfs = [(name, pd.read_excel(xls, sheet_name=name)) for name in xls.sheet_names]
    except Exception as e:
        return jsonify({'success': False, 'error': f'读取文件失败：{str(e)}'}), 400

    total_imported = 0
    total_skipped = 0
    sheet_results = []

    for idx, (sheet_name, df) in enumerate(dfs):
        # 确定该 sheet 的科目和列映射
        if sheet_configs and idx < len(sheet_configs):
            sheet_subject = sheet_configs[idx].get('subject', 'politics')
            col_map = sheet_configs[idx].get('col_map', {})
            extra_fields = sheet_configs[idx].get('extra_fields', {})  # {db_field_label: excel_col_name}
        else:
            sheet_subject = request.form.get('subject', 'politics')
            col_map = {}
            extra_fields = {}

        # 如果前端没传映射，用旧的自动匹配逻辑
        if not col_map:
            for col in df.columns:
                cl = str(col).strip().lower()
                if cl in ('科目', 'subject'):
                    col_map['subject'] = col
                elif cl in ('题号', '编号', 'number', 'question_number', 'no', '序号'):
                    col_map['question_number'] = col
                elif cl in ('标题', 'title', '题目标题'):
                    col_map['title'] = col
                elif cl in ('题目', '题目内容', 'content', 'question'):
                    col_map['content'] = col
                elif cl in ('标准答案', '答案', 'standard_answer', 'answer'):
                    col_map['standard_answer'] = col
                elif cl in ('满分', 'max_score', 'score', '分值'):
                    col_map['max_score'] = col
                elif cl in ('得分点', '评分要点', 'rubric_points', 'points'):
                    col_map['rubric_points'] = col
                elif cl in ('评分规则', 'rubric_rules', 'rules'):
                    col_map['rubric_rules'] = col
                elif cl in ('难度', 'difficulty', '难易'):
                    col_map['difficulty'] = col
                elif cl in ('试卷', '试卷名称', 'exam', 'exam_name', '考试'):
                    col_map['exam_name'] = col

        if 'content' not in col_map:
            sheet_results.append({'sheet': sheet_name, 'imported': 0, 'skipped': 0, 'error': '未找到「题目/题目内容/content」列'})
            continue

        imported = 0
        skipped = 0

        for _, row in df.iterrows():
            content = str(row.get(col_map['content'], '')).strip()
            if not content or content == 'nan':
                skipped += 1
                continue

            title = str(row.get(col_map.get('title', ''), '')).strip()
            if title == 'nan':
                title = ''
            if not title:
                title = content[:30]

            standard_answer = str(row.get(col_map.get('standard_answer', ''), '')).strip()
            if standard_answer == 'nan':
                standard_answer = ''

            max_score_val = row.get(col_map.get('max_score', ''), 10)
            try:
                max_score_val = float(max_score_val) if max_score_val and str(max_score_val) != 'nan' else 10.0
            except:
                max_score_val = 10.0

            row_subject = str(row.get(col_map.get('subject', ''), sheet_subject)).strip()
            if row_subject == 'nan' or not row_subject:
                row_subject = sheet_subject
            # 非 admin 强制用 session.subject
            session_subj = _session_subject()
            if session_subj:
                row_subject = session_subj

            rubric_points = str(row.get(col_map.get('rubric_points', ''), '')).strip()
            if rubric_points == 'nan':
                rubric_points = ''
            rubric_rules = str(row.get(col_map.get('rubric_rules', ''), '')).strip()
            if rubric_rules == 'nan':
                rubric_rules = ''

            # 题号：从 Excel 列中读取
            q_num = str(row.get(col_map.get('question_number', ''), '')).strip()
            if q_num == 'nan' or not q_num:
                q_num = None

            # 难度
            diff_val = str(row.get(col_map.get('difficulty', ''), '')).strip()
            if diff_val == 'nan' or not diff_val:
                diff_val = None

            # 试卷名称
            exam_val = str(row.get(col_map.get('exam_name', ''), '')).strip()
            if exam_val == 'nan' or not exam_val:
                exam_val = None

            rubric_data = {
                'contentType': '简答题',
                'points': [],
            }
            # 存储额外字段到 rubric JSON
            for label, excel_col in extra_fields.items():
                val = str(row.get(excel_col, '')).strip()
                if val and val != 'nan':
                    rubric_data[f'extra_{label}'] = val

            rubric = json.dumps(rubric_data, ensure_ascii=False)

            add_question(
                subject=row_subject,
                title=title,
                content=content,
                original_text=content,
                standard_answer=standard_answer,
                rubric_rules=rubric_rules,
                rubric_points=rubric_points,
                rubric_script='',
                rubric=rubric,
                max_score=max_score_val,
                question_number=q_num,
                difficulty=diff_val,
                exam_name=exam_val
            )
            imported += 1

        total_imported += imported
        total_skipped += skipped
        sheet_results.append({'sheet': sheet_name, 'imported': imported, 'skipped': skipped})

    return jsonify({'success': True, 'data': {'imported': total_imported, 'skipped': total_skipped, 'sheets': sheet_results}})


# Word 导入暂存（token → 解析结果）
_word_import_cache = {}


@api_bp.route('/import-word', methods=['POST'])
def import_word_preview():
    """解析 Word 文档，返回预览（不写入数据库）"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': '未上传文件'}), 400

    file = request.files['file']
    filename = file.filename or ''
    if not filename.lower().endswith('.docx'):
        return jsonify({'success': False, 'error': '仅支持 .docx 格式'}), 400

    try:
        from app.word_parser import parse_english_docx
        questions = parse_english_docx(file.stream)

        if not questions:
            return jsonify({'success': False, 'error': '未解析到题目，请确认文件格式正确（包含【000x】标记）'}), 400

        # 生成 token，完整数据存缓存
        token = f'word_{int(__import__("time").time() * 1000)}'
        _word_import_cache[token] = questions

        # 返回预览数据（不包含大文本字段的完整内容）
        preview = []
        for q in questions:
            preview.append({
                'question_number': q['question_number'],
                'title': q['title'],
                'subject': q['subject'],
                'exam_name': q['exam_name'],
                'max_score': q['max_score'],
                'content_preview': q['content'][:200] + ('...' if len(q['content']) > 200 else ''),
                'standard_answer_preview': q['standard_answer'][:100] + ('...' if len(q['standard_answer']) > 100 else ''),
                'content_length': len(q['content']),
                'rubric_length': len(q['rubric_script']),
            })

        return jsonify({'success': True, 'data': {
            'token': token,
            'filename': filename,
            'total': len(questions),
            'questions': preview,
        }})
    except Exception as e:
        logger.error(f'Word 解析失败: {e}')
        return jsonify({'success': False, 'error': f'解析失败: {str(e)}'}), 500


@api_bp.route('/import-word/confirm', methods=['POST'])
def import_word_confirm():
    """确认导入：通过 token 从缓存获取完整数据，写入数据库"""
    data = request.json
    token = data.get('token', '')
    questions = _word_import_cache.get(token)
    if not questions:
        return jsonify({'success': False, 'error': '导入数据已过期，请重新上传文件'}), 400

    imported = 0
    errors = []
    # 非 admin 强制用 session.subject
    session_subj = _session_subject()
    for q in questions:
        try:
            # 构造 rubric JSON
            rubric = json.dumps({
                'contentType': '简答题',
                'knowledge': '',
                'standardAnswer': q.get('standard_answer', ''),
            }, ensure_ascii=False)

            add_question(
                subject=session_subj or q.get('subject', 'english'),
                title=q.get('title', '') or q.get('content', '')[:30],
                content=q.get('content', ''),
                original_text=q.get('original_text', q.get('content', '')),
                standard_answer=q.get('standard_answer'),
                rubric_rules=q.get('rubric_script', ''),
                rubric_points=q.get('rubric_points', ''),
                rubric_script=q.get('rubric_script', ''),
                rubric=rubric,
                max_score=q.get('max_score', 6),
                question_number=q.get('question_number'),
                difficulty=q.get('difficulty', 'medium'),
                exam_name=q.get('exam_name'),
                content_html=q.get('content_html'),
            )
            imported += 1
        except Exception as e:
            errors.append(f"题号 {q.get('question_number', '?')}: {str(e)}")

    # 清除缓存
    _word_import_cache.pop(token, None)

    return jsonify({'success': True, 'data': {
        'imported': imported,
        'total': len(questions),
        'errors': errors,
    }})
